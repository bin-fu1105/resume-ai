from pathlib import Path

import fitz
from docx import Document


class ResumeParseError(Exception):
    """Raised when resume text cannot be extracted."""


def extract_pdf_text(file_path: str) -> str:
    text_parts = []

    with fitz.open(file_path) as document:
        for page in document:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text.strip())

    return "\n".join(text_parts).strip()


def extract_docx_text(file_path: str) -> str:
    document = Document(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_resume_text(file_path: str) -> str:
    path = Path(file_path)
    extension = path.suffix.lower()

    try:
        if extension == ".pdf":
            text = extract_pdf_text(str(path))
        elif extension == ".docx":
            text = extract_docx_text(str(path))
        else:
            raise ResumeParseError(
                "Unsupported file type for parsing. Only PDF and DOCX are supported."
            )
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError(f"Failed to extract resume text: {exc}") from exc

    if not text:
        raise ResumeParseError("No readable text found in the uploaded resume.")

    return text
