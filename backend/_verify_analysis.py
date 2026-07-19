"""Smoke-test resume parsing and /analyze without printing secrets."""

from pathlib import Path
from urllib import request, error
import json

from docx import Document
import fitz

from services.resume_parser import extract_resume_text


def make_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "Jane Doe\nSoftware Engineer\nSkills: Python, FastAPI, React, SQL\n"
        "Experience: Built APIs and dashboards for SaaS products.",
    )
    doc.save(path)
    doc.close()


def make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Jane Doe", level=1)
    document.add_paragraph("Software Engineer")
    document.add_paragraph("Skills: Python, FastAPI, React, SQL")
    document.add_paragraph(
        "Experience: Built APIs and dashboards for SaaS products."
    )
    document.save(path)


def upload(path: Path):
    data = path.read_bytes()
    boundary = "----BoundaryVerifyAnalysis"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{path.name}"\r\n'
        f"Content-Type: application/octet-stream\r\n\r\n"
    ).encode() + data + f"\r\n--{boundary}--\r\n".encode()
    req = request.Request(
        "http://127.0.0.1:8000/upload",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with request.urlopen(req) as resp:
        return json.loads(resp.read().decode())


def analyze(filename: str, job_description: str):
    payload = json.dumps(
        {
            "filename": filename,
            "job_description": job_description,
        }
    ).encode()
    req = request.Request(
        "http://127.0.0.1:8000/analyze",
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with request.urlopen(req, timeout=120) as resp:
            return resp.status, json.loads(resp.read().decode())
    except error.HTTPError as exc:
        return exc.code, json.loads(exc.read().decode())


pdf_path = Path("_verify_resume.pdf")
docx_path = Path("_verify_resume.docx")
make_pdf(pdf_path)
make_docx(docx_path)

pdf_text = extract_resume_text(str(pdf_path))
docx_text = extract_resume_text(str(docx_path))
print("PDF parse OK:", "Python" in pdf_text and len(pdf_text) > 20)
print("DOCX parse OK:", "Python" in docx_text and len(docx_text) > 20)

upload_pdf = upload(pdf_path)
print("Upload PDF:", upload_pdf.get("status"), bool(upload_pdf.get("filename")))

status, result = analyze(
    upload_pdf["filename"],
    "Looking for a Software Engineer with Python, FastAPI, React, and SQL experience.",
)
print("Analyze status:", status)
analysis = result.get("analysis") or {}
required = [
    "ats_score",
    "ats_explanation",
    "resume_match",
    "missing_skills",
    "strengths",
    "suggestions",
    "optimized_summary",
]
print("Has all keys:", all(key in analysis for key in required))
if status == 200:
    print("ats_score:", analysis.get("ats_score"))
    print("has ats_explanation:", bool(analysis.get("ats_explanation")))
    print("resume_match keys:", sorted((analysis.get("resume_match") or {}).keys()))
    print("missing_skills count:", len(analysis.get("missing_skills") or []))
    print("strengths count:", len(analysis.get("strengths") or []))
    print("suggestions count:", len(analysis.get("suggestions") or []))
    first_suggestion = (analysis.get("suggestions") or [{}])[0]
    print(
        "suggestion has reason/example/impact:",
        isinstance(first_suggestion, dict)
        and {"reason", "example", "impact"} <= set(first_suggestion.keys()),
    )
    print("summary length:", len(analysis.get("optimized_summary") or ""))
else:
    print("Analyze error detail:", result.get("detail") or result)

pdf_path.unlink(missing_ok=True)
docx_path.unlink(missing_ok=True)
